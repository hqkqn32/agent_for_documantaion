
from docx import Document
from datetime import datetime
import os

class WordExporter:
    def __init__(self, export_dir="exports"):
        self.export_dir = export_dir
        os.makedirs(self.export_dir, exist_ok=True)

    def export(self, title, content):
        filename = f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.export_dir, filename)

        doc = Document()
        doc.add_heading(title, 0)

        for paragraph in content.strip().split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        doc.save(filepath)
        print(f"\n Word dosyası oluşturuldu: {filepath}")
        return filepath

# Test
if __name__ == "__main__":
    exporter = WordExporter()
    title = "Test Dökümantasyon"
    content = """
    Bu belge test amaçlı oluşturulmuştur.

    - Amaç: Word dışa aktarma işlemini denemek
    - Yöntem: python-docx kütüphanesi
    """
    exporter.export(title, content)
