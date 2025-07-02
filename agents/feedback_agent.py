

from docx import Document
from datetime import datetime
import os

class WordExporter:
    def __init__(self, export_dir="exports"):
        self.export_dir = export_dir
        os.makedirs(self.export_dir, exist_ok=True)

    def export(self, title, content):
        # BaÅŸlÄ±k, dosya adÄ±na uygun hale getirilir
        filename = f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.export_dir, filename)

        doc = Document()
        doc.add_heading(title, 0)

        for paragraph in content.strip().split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        doc.save(filepath)
        print(f"\nğŸ“„ Word dosyasÄ± oluÅŸturuldu: {filepath}")
        return filepath

# Test
if __name__ == "__main__":
    exporter = WordExporter()
    title = "Test DÃ¶kÃ¼mantasyon"
    content = """
    Bu belge test amaÃ§lÄ± oluÅŸturulmuÅŸtur.

    - AmaÃ§: Word dÄ±ÅŸa aktarma iÅŸlemini denemek
    - YÃ¶ntem: python-docx kÃ¼tÃ¼phanesi
    """
    exporter.export(title, content)
